from flask import Flask, render_template, request, send_file
import io
from google import genai
from docx import Document
import markdown # Pour formater les exercices & corrections
import os
from dotenv import load_dotenv
load_dotenv()

app = Flask(__name__)

# Configuration API

CLE_API = os.environ.get("GOOGLE_API_KEY")

if not CLE_API:
    print("ERREUR : La clé d'API n'est pas renseignée correctement.")
else:
    client = genai.Client(api_key=CLE_API)



@app.route("/")
def index():
    return render_template("genererForm.html", methods=['POST']) # Render la page HTML formulaire de génération d'exercice

@app.route("/formulaire_generer", methods=['POST'])
def generer():
    matiere = request.form.get('matiere')
    niveau = request.form.get('niveau')
    precisions = request.form.get('precisions')

    prompt = f"""
    #Rôle
    Enseignant en {matiere}, niveau {niveau}
    #Contexte 
    Rédaction d'un contrôle
    #Objectif 
    Créer des exercices complets pour un contrôle en classe.
    Les exercices doivent être d'une difficulté adaptée à un élève de {niveau}.
    Potentielles précisions, directives à suivre : [{precisions}] (ignorer si vide)
    #Livrables
    Des exercices structurés de manière claire :
    - Un titre de sujet.
    - Une introduction contextuelle si nécessaire.
    - Plusieurs questions distinctes (numérotées).
    - Pour chaque question, laisse une zone de texte vide (indiquée par "[Réponse ici]") d'environ 4-5 lignes pour que l'élève puisse répondre.
    #Calibrage / Format
    Proposer entre 5 et 10 exercices de {matiere}
    #Validation / Contraintes
    Les exercices doivent impérativement être au niveau {niveau}
    Utiliser une mise en page classique de contrôle scolaire
    L'exercice doit être prêt à être imprimé et donné à un élève.
    N'inclus aucune note de l'IA ni de texte d'introduction/conclusion. Juste le sujet.
    Utilise une mise en page claire avec des sauts de ligne.    
    """

    try:
        response = client.models.generate_content(
            model="gemini-3-flash-preview",
            contents=prompt
        )
        exercices = response.text
    
    except Exception as e:
        print(f"Erreur lors de la génération du contenu : {e}")
        exercices = "Une erreur est survenue lors de la génération des exercices. Veuillez réessayer plus tard."

    exercices_formates = markdown.markdown(exercices)
    return render_template("generer.html", matiere=matiere, niveau=niveau, precisions=precisions, exercices=exercices, exercices_formates = exercices_formates)

@app.route("/telecharger_exercice", methods=['POST'])
def telecharger():
    exercices = request.form.get('exercices')
    matiere = request.form.get('matiere')
    niveau = request.form.get('niveau')
    document = Document()
    document.add_heading(f"Exercices de {matiere} - Niveau {niveau}", level=0)
    document.add_paragraph(exercices)
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"exercices_{matiere}_{niveau}.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route("/upload_exercice")
def televerser():
    return render_template("corrigerForm.html", methods=['POST'])

@app.route("/corriger", methods=['POST'])
def corriger():
    exercice = request.files.get('exercice')
    if not exercice:
        return "Aucun fichier téléchargé. Veuillez réessayer.", 400
    
    try:
        document = Document((io.BytesIO(exercice.read())))
        liste = []
        for para in document.paragraphs:
            if para.text.strip(): # Si le paragraphe n'est pas vide
                liste.append(para.text) # On ajoute son texte
        contenu = "\n".join(liste) # On concatène tous les textes avec des sauts de ligne

    except Exception as e:
        print(f"Erreur lors de la lecture du fichier : {e}")
        return "Une erreur est survenue lors de la lecture du fichier. Veuillez réessayer avec un fichier valide.", 400
    
    prompt = f"""
    #Rôle
    Correcteur
    #Contexte
    Correction d'un sujet d'exercices
    #Objectif
    Fournir une correction détaillée et complète pour chaque exercice du sujet ainsi qu'une note estimée sur 20, avec des conseils.
    #Livrables
    - Pour chaque exercice, une correction complète et détaillée.
    - Une note globale sur 20 pour l'ensemble du sujet.
    - Des conseils pour améliorer les performances de l'élève.
    #Calibrage / Format
    - La correction doit être claire, structurée et facile à comprendre.
    - La note doit être justifiée par des critères précis.
    - Les conseils doivent être pertinents et adaptés au niveau de l'élève.
    #Validation / Contraintes
    - La correction doit être précise et complète, couvrant tous les aspects des exercices.
    - La correction doit correspondre au niveau attendu pour le sujet d'exercices.
    N'inclus aucune note de l'IA ni de texte d'introduction/conclusion. Juste la correction
    Utilise une mise en page claire avec des sauts de ligne.
    #Sujet d'exercices à corriger
    {contenu}
    """
    
    try:
        response = client.models.generate_content(
            model="gemini-3-flash-preview",
            contents=prompt
        )
        exercices_corriges = response.text
    
    except Exception as e:
        print(f"Erreur lors de la correction du contenu : {e}")
        exercices_corriges = "Une erreur est survenue lors de la correction des exercices. Veuillez réessayer plus tard."

    correction_formatee = markdown.markdown(exercices_corriges)
    return render_template("correction.html", exercices_corriges=exercices_corriges, correction_formatee=correction_formatee)

if __name__ == "__main__":
    app.run(debug=False)


